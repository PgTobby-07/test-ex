"""
Converts SAD_EWP.md to a formatted Word document (SAD_EWP.docx).
Requires: pip install python-docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = Path(r"C:\Users\ridha\.gemini\antigravity\brain\a298afe7-29ed-4248-a503-24ba14a8be1d\SAD_EWP.md")
OUT_PATH = Path(r"C:\Users\ridha\Downloads\SAD_EWP.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(text, level):
    h = doc.add_heading(text.strip(), level=level)
    h.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_code_block(code_text):
    """Adds a grey-background monospace paragraph for each line."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        # light grey shading
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F0F0F0")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

def add_table_from_md(lines):
    """Parse a markdown table and add it as a Word table."""
    rows = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    cols = max(len(r) for r in rows)
    tbl  = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= cols:
                break
            cell = row.cells[c_idx]
            cell.text = cell_text
            # Header row bold
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # blue header background
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd  = OxmlElement("w:shd")
                    shd.set(qn("w:val"),   "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"),  "1F497D")
                    tcPr.append(shd)

    doc.add_paragraph()  # spacing after table

# ── Parser ────────────────────────────────────────────────────────────────────
lines = MD_PATH.read_text(encoding="utf-8").splitlines()

i = 0
while i < len(lines):
    line = lines[i]

    # Horizontal rule
    if re.match(r"^---+$", line.strip()):
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        set_heading(m.group(2), level)
        i += 1
        continue

    # Fenced code block  (``` or ```lang)
    if line.strip().startswith("```"):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        add_code_block("\n".join(code_lines))
        doc.add_paragraph()
        i += 1
        continue

    # Markdown table
    if line.strip().startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        add_table_from_md(table_lines)
        continue

    # Bullet list items
    m = re.match(r"^(\s*)[-*]\s+(.*)", line)
    if m:
        indent = len(m.group(1)) // 2
        style  = "List Bullet 2" if indent > 0 else "List Bullet"
        p = doc.add_paragraph(style=style)
        # inline bold **text**
        text = m.group(2)
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # Blank line
    if not line.strip():
        doc.add_paragraph()
        i += 1
        continue

    # Normal paragraph (with inline bold/code)
    p = doc.add_paragraph(style="Normal")
    text = line
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:
            p.add_run(part)
    i += 1

doc.save(OUT_PATH)
print(f"Done! Saved to: {OUT_PATH}")
